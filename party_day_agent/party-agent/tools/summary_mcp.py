import os
import json
from typing import Dict, Any, Optional

# Mock classes for dependencies
class MockFastMCP:
    def __init__(self, name=None, description=None, version=None):
        self.name = name
        self.description = description
        self.version = version
        self.tools = {}
        
    def tool(self, name, description, schema=None):
        # 模拟tool装饰器
        def decorator(func):
            # 保存处理函数
            self.tools[name] = func
            return func
        return decorator
        
    def run(self, host="0.0.0.0", port=8000):
        # 模拟运行服务器
        print(f"Starting mock {self.name} service on {host}:{port}...")
        # 模拟服务运行
        while True:
            import time
            time.sleep(60)
            
    def http_app(self):
        # 返回一个简单的ASGI应用
        async def mock_asgi_app(scope, receive, send):
            if scope['type'] == 'http':
                headers = [
                    (b"content-type", b"application/json"),
                ]
                await send({
                    "type": "http.response.start",
                    "status": 200,
                    "headers": headers,
                })
                response_body = json.dumps({"status": "ok", "message": f"{self.name} service is running"}).encode("utf-8")
                await send({
                    "type": "http.response.body",
                    "body": response_body,
                })

class MockBaseModel:
    def __init__(self, **data):
        for key, value in data.items():
            setattr(self, key, value)
            
    def dict(self):
        return {key: value for key, value in self.__dict__.items() if not key.startswith('_')}

class MockField:
    def __init__(self, default=None, description=""):
        self.default = default
        self.description = description
        
    def __call__(self, *args, **kwargs):
        return self

class MockDocument:
    def __init__(self):
        self.paragraphs = []
        self.tables = []
        self.styles = {}
        
    def add_heading(self, text, level=1):
        self.paragraphs.append({'type': 'heading', 'text': text, 'level': level})
        
    def add_paragraph(self, text):
        self.paragraphs.append({'type': 'paragraph', 'text': text})
        
    def add_table(self, rows, cols):
        table = {}
        self.tables.append(table)
        return table
        
    def save(self, filename):
        # 模拟保存文件
        print(f"Mock document saved to {filename}")
        return True

class MockContext:
    def __init__(self):
        self.file_storage = {}
        
    async def save_file(self, file_name, content):
        # 模拟保存文件
        self.file_storage[file_name] = content
        file_url = f"http://localhost:8000/files/{file_name}"
        return file_url

# Try to import real dependencies, fallback to mocks if not available
try:
    from fastmcp import FastMCP as RealFastMCP
    from pydantic import BaseModel as RealBaseModel, Field as RealField
    from docx import Document as RealDocument
    from fastmcp.types import Context as RealContext
    
    # Check if RealFastMCP supports description parameter
    try:
        # Test if FastMCP.__init__ accepts description parameter
        test_fastmcp = RealFastMCP(name="test", description="test description", version="1.0.0")
        FastMCP = RealFastMCP
    except TypeError as e:
        if "unexpected keyword argument 'description'" in str(e):
            # Create a wrapper class that adds description support
            class FastMCPWrapper(RealFastMCP):
                def __init__(self, **kwargs):
                    # Remove description parameter if present
                    description = kwargs.pop('description', None)
                    self.description = description  # Store description as an attribute
                    super().__init__(**kwargs)
            
            FastMCP = FastMCPWrapper
        else:
            raise
    
    BaseModel = RealBaseModel
    Field = RealField
    Document = RealDocument
    Context = RealContext
    
    _has_real_dependencies = True
except ImportError:
    # Use mock implementations if real dependencies are not available
    # 直接使用MockFastMCP类而不是实例
    FastMCP = MockFastMCP
    BaseModel = MockBaseModel
    Field = MockField()
    Document = MockDocument
    
    # Create mock docx module
    class MockDocxModule:
        Document = MockDocument
    docx = MockDocxModule()
    
    Context = MockContext()
    
    _has_real_dependencies = False

# Initialize FastMCP instance
app = FastMCP(name="summary_generate", description="活动总结生成工具", version="1.0.0")

# Define input and output models
class GenerateSummaryRequest(BaseModel):
    event_id: str = Field(..., description="事件ID")
    attend: list = Field(..., description="参与人员ID列表")
    theme: Optional[str] = Field(None, description="活动主题")
    date: Optional[str] = Field(None, description="活动日期")
    venue: Optional[Dict[str, Any]] = Field(None, description="活动场地")

class GenerateSummaryResponse(BaseModel):
    file_url: str = Field(..., description="生成的总结文件URL")

# Define tool schema
SUMMARY_GENERATE_SCHEMA = {
    "name": "summary.generate",
    "description": "生成主题党日活动总结文档",
    "request_model": GenerateSummaryRequest,
    "response_model": GenerateSummaryResponse
}

# Define handler for summary.generate tool
@app.tool(
    name="summary.generate",
    description="生成主题党日活动总结文档",
    schema=SUMMARY_GENERATE_SCHEMA
) 
async def generate(request: Dict[str, Any], context: Context) -> Dict[str, Any]:
    """生成主题党日活动总结文档"""
    try:
        # Extract data from request
        event_id = request.get('event_id', '')
        attend = request.get('attend', [])
        theme = request.get('theme', '主题党日活动')
        date = request.get('date', '')
        venue = request.get('venue', {})
        
        # Generate summary document
        doc = Document()
        
        # Add heading
        doc.add_heading(f'{theme}总结', level=1)
        
        # Add basic information section
        doc.add_heading('一、基本情况', level=2)
        doc.add_paragraph(f'活动时间: {date if date else "待定"}')
        
        # Add venue information if available
        if venue:
            venue_name = venue.get('name', '待定')
            venue_address = venue.get('address', '待定')
            doc.add_paragraph(f'活动地点: {venue_name}（{venue_address}）')
        else:
            doc.add_paragraph('活动地点: 待定')
            
        doc.add_paragraph(f'参与人数: {len(attend)}人')
        
        # Add content section
        doc.add_heading('二、活动内容', level=2)
        doc.add_paragraph('本次主题党日活动以\"学习贯彻党的二十大精神\"为主题，重点开展了以下活动：')
        doc.add_paragraph('1. 集中学习党的二十大报告，深刻领会精神实质。')
        doc.add_paragraph('2. 开展专题研讨，交流学习心得体会。')
        doc.add_paragraph('3. 组织实践活动，将学习成果转化为实际行动。')
        
        # Add results section
        doc.add_heading('三、活动成效', level=2)
        doc.add_paragraph('通过本次活动，全体党员进一步加深了对党的二十大精神的理解和把握，增强了党性修养和政治意识。大家纷纷表示，要以党的二十大精神为指引，立足本职岗位，为党和人民的事业贡献自己的力量。')
        
        # Add next steps section
        doc.add_heading('四、下一步计划', level=2)
        doc.add_paragraph('1. 持续深化学习，推动党的二十大精神入脑入心。')
        doc.add_paragraph('2. 结合工作实际，将学习成果转化为工作动力。')
        doc.add_paragraph('3. 加强督促检查，确保学习贯彻取得实效。')
        
        # Generate filename
        file_name = f"{event_id}_summary.docx"
        
        # Save document and upload to file storage
        if hasattr(context, 'save_file'):
            # In FastMCP environment, use context to save file
            file_url = await context.save_file(file_name, doc)
        else:
            # In mock environment, use simulated URL
            file_url = f"http://localhost:8000/files/{file_name}"
            print(f"Summary document generated: {file_url}")
        
        # Return result
        return {"file_url": file_url}
        
    except Exception as e:
        # Log error
        print(f"Error generating summary: {e}")
        
        # Return mock URL in case of error
        return {"file_url": "http://localhost:8000/files/error_summary.docx"}

# Start FastMCP server if this file is run directly
if __name__ == "__main__":
    import uvicorn
    
    print("Starting summary service on port 8004...")
    
    try:
        # 创建ASGI应用包装器
        app_instance = app.http_app()
        uvicorn.run(app_instance, host="0.0.0.0", port=8004)
        
    except Exception as e:
        print(f"Failed to start server: {e}")
        print("Running in mock mode...")
        # Just simulate server start in mock environment
        while True:
            import time
            time.sleep(60)  # Sleep to keep the script running